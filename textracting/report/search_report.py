## @file
# Main file for identifying/classifying parts of a report and bookmarking it, drawing on it, and saving its sections.
# by Anna Andraszek

import json
import pandas as pd
import numpy as np
from report import toc_classification, heading_id_toc, marginals_classification, page_extraction, fig_classification, heading_id_intext, heading_classification
import paths
import os
os.environ['KMP_AFFINITY'] = 'noverbose'  # this stops a lot of warning messages being printed
from pdf2image import convert_from_path
from PIL import ImageDraw, Image
import re
from PyPDF2 import PdfFileWriter, PdfFileReader
import time
import docx
from report.heading_id_intext import Text2CNBPrediction, Num2Cyfra1
#import textdistance
import jsonpickle
import glob

mode = paths.production
#mode = paths.dataset_version


class Report():
    def __init__(self, docid: str, filenum: str):
        ## unique report identifier
        self.docid = docid
        ## desired file number of that report
        self.filenum = filenum
        #self.toc_dataset_path = settings.production_path + docid + '_toc_dataset.csv'
        ## document information: contents of the restructpageinfo file
        self.docinfo = self.get_doc_info()
        if len(self.docinfo.keys()) == 0:
            print("ERROR: docinfo is empty. Cannot process Report", docid)
            raise ValueError
        ## lines making up the text of the document
        self.doclines = self.get_doc_lines()
        ## dataset of lines for ML models
        self.line_dataset = self.create_line_dataset()
        ## Table of Contents page number
        self.toc_page = self.get_toc_page()
        ## Page numbers of pages which majorly contain figures/images
        self.fig_pages = self.get_fig_pages()  #fig_classification.get_fig_pages(self.docid, self.docinfo, self.doclines)

        if self.toc_page:
            self.headings, self.subheadings = self.get_headings()
        else:
            self.headings, self.subheadings = pd.DataFrame(), pd.DataFrame()
        self.get_marginals()

        ## List of dictionaries of HeadingText, PageNum, and LineNum attributes pointing to the beginnings of sections in the document
        self.section_ptrs = self.get_section_ptrs()  # section_ptrs = [{HeadingText: , PageNum: , LineNum: }]
        ## Sectioned document text
        self.section_content = self.get_sections()
        #self.toc_heading_classes, self.text_heading_classes = self.classify_headings()

    ## Finds pages majorly containing figures/images by using model in fig_classification
    def get_fig_pages(self):
        data = fig_classification.create_individual_dataset(self.docid, self.docinfo, self.doclines)
        fig_classification.get_fig_pages(data, mode)

    ## Finds headers/footers
    def get_marginals(self):
        self.marginals = marginals_classification.get_marginals(
            self.create_marginals_dataset(), mode)  # a df containing many columns, key: pagenum, text
        self.marginals_set = set([(p, l) for p, l in zip(self.marginals.PageNum, self.marginals.LineNum)])
        #self.page_nums = page_extraction.get_page_nums(self.marginals, mode=mode)
        self.page_nums = None

    ## Matches TOC headings to in-text headings
    def match_headings(self, headings_intext):
        all_headings = pd.concat((self.headings, self.subheadings))
        # there can only be one of each value in self.headings_intext.MatchesI
        matched_is = []
        for value in headings_intext.MatchesI.values:
            if value not in matched_is:
                matched_is.append(value)
            else:
                temp = headings_intext.loc[headings_intext['MatchesI'] == value]
                imax = temp['MatchesHeading'].idxmax()
                save = temp.loc[imax]
                headings_intext = headings_intext.drop(index=temp.index.values)  # remove all non matching rows so this only has to be done once per i
                headings_intext = headings_intext.append(save)  # don't need to remove and re-append i because it's just a number
        # sort headings
        headings_intext.sort_values(by=['PageNum', 'LineNum'], inplace=True)
        return headings_intext

    ## Finds starting lines of sections in the document
    def get_section_ptrs(self):
        headings_intext = heading_id_intext.get_headings_intext(self.create_intext_id_dataset(), self.toc_page, mode)
        # compare intext headings vs toc headings and choose only 1:1
        self.headings_intext = self.match_headings(headings_intext)
        section_ptrs = self.headings_intext.loc[self.headings_intext['Heading'] == 1]
        self.subsection_ptrs = self.headings_intext.loc[self.headings_intext['Heading'] == 2]
        self.subsection_ptrs.reset_index(inplace=True, drop=True)
        section_ptrs.reset_index(inplace=True, drop=True)
        return section_ptrs

    ## UNUSED Classifies headings into major categories
    def classify_headings(self):
        # classify both toc headings and text headings, but separately
        intext_dataset = self.headings_intext
        if (self.toc_page) and (self.headings.shape[0] > 0):
            toc_dataset = self.line_dataset.loc[self.line_dataset['PageNum'] == self.toc_page]
            dataset = toc_dataset[toc_dataset['LineNum'].isin(self.headings['LineNum'])]
            if self.subheadings.shape[0] > 0:
                toc_head_dataset = dataset.append(toc_dataset[toc_dataset['LineNum'].isin(self.subheadings['LineNum'])], ignore_index=True)
            else:
                toc_head_dataset = dataset
            toc_res = heading_classification.classify(toc_head_dataset)
            #print(toc_res)
        else: toc_res = []
        if intext_dataset.shape[0] > 0:
            intext_res = heading_classification.classify(intext_dataset)
            #print(intext_res)
        else: intext_res = []
        return toc_res, intext_res

    ## Gets lines of text in the document
    def get_doc_lines(self):
        pagelines = {}
        prev_pg = 0
        for page in self.docinfo.items():
            pagenum = page[0]
            if int(pagenum) > prev_pg + 1:  # if there are empty pages, still add blank dict entries to avoid messing up len
                empty_pagenum = prev_pg + 1
                pagelines[str(empty_pagenum)] = []
                while empty_pagenum < int(pagenum):
                    empty_pagenum += 1
                    pagelines[str(empty_pagenum)] = []
            info = page[1]
            lines = []
            for line in info:
                lines.append(line['Text'])
            pagelines[pagenum] = lines
            prev_pg = int(pagenum)
        return pagelines

    ## Reads the restructpageinfo file into memory
    def get_doc_info(self):
        pageinfo = paths.get_restructpageinfo_file(self.docid, file_num=self.filenum)
        pi = json.load(open(pageinfo, "r"))
        return pi

    ## Creates a dataset for classifying a page as a table of contents (TOC)
    def create_toc_dataset(self):
        include_cols = toc_classification.include_cols
        if mode == paths.production:
            include_cols = include_cols[:-2]
        docset = np.zeros((len(self.docinfo.items()), len(include_cols)+1))
        for info, lines, j, in zip(self.docinfo.items(), self.doclines.items(), range(len(self.docinfo.items()))):
            toc = 0
            c = 0
            listof = 0
            for line in lines[1]:
                if 'contents' in line.lower():
                    c = 1
                    if 'table of contents' in line.lower():
                        toc = 1
                if 'list of' in line.lower():
                    listof = 1
            if mode == paths.production:
                row = [self.docid, info[0], len(lines[1]), toc, c]
            else:
                row = [self.docid, info[0], len(lines[1]), toc, c, listof, 0]
            docset[j] = np.array(row)
        columns = ['DocID']
        columns.extend(include_cols)
        pgdf = pd.DataFrame(data=docset, columns=columns)
        #pgdf.to_csv(self.toc_dataset_path, index=False)
        return pgdf

    ## Gets TOC page number of the document (if exists)
    def get_toc_page(self):
        data = self.create_toc_dataset()
        toc_pages = toc_classification.get_toc_pages(data, mode)
        try:
            toc_pages = toc_pages.reset_index()
            tc = toc_pages.iloc[toc_pages['proba'].idxmax()]  # getting most probable toc page
            toc = int(tc['PageNum']) # how do you account for multiple toc pages?
        except (IndexError, ValueError):
            print("TOC page doesn't exist")
            toc = None
        return toc

    ## Get section headings from the TOC
    def get_headings(self):
        df = self.create_identification_dataset()
        #df = heading_id_toc.pre_process_id_dataset(datafile=df, training=False)
        preds = heading_id_toc.get_toc_headings(df, mode)
        headings = preds.loc[preds['Heading'] == 1]
        subheadings = preds.loc[preds['Heading'] == 2]
        return headings, subheadings

    ## Creates dataset of document lines, to be used by ML methods
    def create_line_dataset(self):
        # create a dataset of all lines in the document and their universally used attributes like words2width, centrality
            # to be used by any classifier which uses lines
        columns = ['DocID', 'PageNum', 'LineNum', 'NormedLineNum', 'Text', 'Words2Width', 'WordsWidth', 'Width',
                   'Height', 'Left', 'Top', 'Centrality', 'WordCount']
        df = pd.DataFrame(columns=columns)
        for info in self.docinfo.items():
            docset = []
            page = info[0]
            for line in info[1]:
                bb = line['BoundingBox']
                centrality = 0.5 - abs(bb['Left'] + (bb['Width'] / 2) - 0.5)  # the higher value the more central
                words2width = line['WordsWidth'] / bb['Width']
                wordcount = len(line['Text'].split())
                docset.append([self.docid, int(page), line['LineNum'], 0, line['Text'], words2width, line['WordsWidth'],
                               bb['Width'], bb['Height'], bb['Left'], bb['Top'], centrality, wordcount])
            temp = pd.DataFrame(data=docset, columns=columns)
            if (max(temp['LineNum']) - min(temp['LineNum'])) == 0:  # only one line # avoid NaN from div by 0
                temp['NormedLineNum'] = 0
            else:
                temp['NormedLineNum'] = (temp['LineNum'] - min(temp['LineNum'])) / (
                        max(temp['LineNum']) - min(temp['LineNum']))

            df = df.append(temp, ignore_index=True)
        unnormed = np.array(df['Centrality'])
        normalized = (unnormed - min(unnormed)) / (max(unnormed) - min(unnormed))
        df['Centrality'] = normalized
        return df

    ## Creates dataset for identifying headings in TOC
    def create_identification_dataset(self):
        columns = ['DocID', 'LineNum','Left', 'Top', 'Text']#, 'Width', 'Height', ]
        df = self.line_dataset.loc[self.line_dataset['PageNum'] == self.toc_page]
        df = df[columns]
        df.reset_index(inplace=True, drop=True)
        return df

    ## Creates dataset for identifying headings in text
    def create_intext_id_dataset(self):
        #df = self.line_dataset.copy(deep=True)
        df = self.line_dataset.loc[self.line_dataset['PageNum'] != self.toc_page]
        df['ContainsNum'] = df.Text.apply(lambda x: heading_id_intext.contains_num(x))
        if self.headings.shape[0] > 0:
            #self.headings['Text'] = self.headings.apply(lambda x: x.SectionPrefix + ' ' + x.SectionText, axis=1)
            self.headings['Heading'] = 1
        if self.subheadings.shape[0] > 0:
            #self.subheadings['Text'] = self.subheadings.apply(lambda x: x.SectionPrefix + ' ' + x.SectionText, axis=1)
            self.subheadings['Heading'] = 2
        headings = pd.concat([self.headings, self.subheadings])#, ignore_index=True)
        df['MatchesHeading'], df['MatchesType'], df['MatchesI'] = heading_id_intext.compare_lines2headings(df.Text.values, headings)
        df['Heading'] = None
        return df

    ## Creates dataset of marginals lines, to be used by ML methods
    def create_marginals_dataset(self):
        df = self.line_dataset.copy(deep=True)
        df['ContainsNum'] = df.Text.apply(lambda x: marginals_classification.contains_num(x))
        df['ContainsTab'] = df.Text.apply(lambda x: marginals_classification.contains_tab(x))
        df['ContainsPage'] = df.Text.apply(lambda x: marginals_classification.contains_page(x))
        #df['Marginal'] = 0  # no y column
        df.drop(columns=['WordCount'], inplace=True)
        return df

    ## Gets text of document sections
    def get_sections(self):
        # from section ptrs, section = section ptr, reading until the start of the next section
        if self.section_ptrs.shape[0] == 0:
            return []

        section_num = 0
        ptr = self.section_ptrs.iloc[section_num]
        name = ptr['Text']
        start_page = ptr['PageNum']
        start_line = ptr['LineNum']
        if self.section_ptrs.shape[0] > 1:
            next_section = self.section_ptrs.iloc[section_num+1]
            end_page = next_section['PageNum']
            end_line = next_section['LineNum']
        else:
            end_page = len(self.doclines)
            end_line = len(self.doclines[str(end_page)])
        content = []
        sections = []
        end = False
        for page in range(start_page, len(self.doclines.items()) +1): # +1 because index starts at 1, +1 to include last element
            #try:
            for linenum in range(len(self.doclines[str(page)])):
                if (page, linenum) not in self.marginals_set:  # if line is not a marginal
                    line = self.doclines[str(page)][linenum]
                    if page == end_page and linenum == end_line:  # if end of section
                        section = {'Heading': name, 'Content': content}
                        sections.append(section)

                        if section_num != len(self.section_ptrs) - 1: # if there is a next section
                            content = []
                            content.append(line)
                            section_num +=1
                            ptr = self.section_ptrs.iloc[section_num]
                            name = ptr['Text']

                        if section_num != len(self.section_ptrs) - 1: # if next section is not last section
                            next_section = self.section_ptrs.iloc[section_num + 1]
                            end_page = next_section['PageNum']
                            end_line = next_section['LineNum']

                        else:   # if next section is last section
                            end_page = len(self.doclines.items())
                            end_line = len(self.doclines[str(end_page)])
                            end = True

                    elif end and page == end_page and linenum == end_line-1: # if current section is last section, current line is last line in whole doc
                        content.append(line)
                        section = {'Heading': name, 'Content': content}
                        sections.append(section)

                    elif page == start_page and linenum < start_line:  # if before start line on start page
                        continue

                    elif page == end_page and linenum == end_line-1:  # stop the heading line being added to the end of the previous section
                        continue

                    else:  # add line to content
                        #line = self.doclines[str(page)][linenum]
                        content.append(line)
            #except KeyError:
            #    print('Page ' + str(page) + ' is missing')
        return sections

## Prints out text of document sections
def print_sections(report):
    print("Sections at: \n", report.section_ptrs)
    # json.dump(r.section_content, open('sections.json', 'w'))
    for section in report.section_content:
        print(section['Heading'])
        for line in section['Content']:
            print(line)
        print('\n')


## Draws boundings boxes which higlight various elements identified in report and saves this to a new file.
# Used for visual confirmation of identification effectiveness only.
def draw_report(report):
    report_path = paths.get_report_name(report.docid, local_path=True, file_extension='.pdf', file_num=report.filenum)
    images = convert_from_path(report_path)

    doc = report.docinfo
    drawn_images = []

    for page in doc.items():
        i = int(page[0])-1
        image = images[i]  # this has to be of type RGB
        width, height = image.size
        draw = ImageDraw.Draw(image, 'RGBA')

        if int(page[0]) in report.marginals['PageNum'].values:     # draw bb around marginals
            lnnum = report.marginals.loc[report.marginals['PageNum'] == int(page[0])]['LineNum']
            for ln in lnnum.values:
                linenum = ln - 1
                line = page[1][linenum]
                box = line['BoundingBox']
                left = width * box['Left']
                top = height * box['Top']
                draw.rectangle([left, top, left + (width * box['Width']), top + (height * box['Height'])], outline='orange')

            # draw bb around page number (by comparing marginal content to result of page number extraction)
            if isinstance(report.page_nums, pd.DataFrame):
                if int(page[0]) in report.page_nums['PageNum'].values:  # draw bb around marginals
                    pg_marginal = report.page_nums.loc[report.page_nums['PageNum'] == int(page[0])]
                    #pglnnum = pg_marginal['LineNum']
                    #pglinenum = pglnnum.values[0] - 1
                    text = pg_marginal.Text.values[0]
                    split_text = text.split('\t')
                    reg = r'(^|\s)' + str(pg_marginal['Page'].values[0]) + r'($|\s)'   # implement returning pagenum position instead? would make this MUCH easier
                    pgnum_i = None
                    for t, i in zip(split_text, range(len(split_text))):
                        if re.search(reg, t):
                            pgnum_i = i
                            break
                    if pgnum_i:
                        box = line['OriginalBBs'][pgnum_i]
                        left = width * box['Left']
                        top = height * box['Top']
                        draw.rectangle([left, top, left + (width * box['Width']), top + (height * box['Height'])],
                                   outline='red')

                #original_marginal_bb = docinfo[pagestr][lineindex]['OriginalBBs'][index in marginal]


        if page[0] == str(report.toc_page): # change colour of toc page
            # for i, row in report.toc_dataset.iterrows():  # did this mean to put rectangles around toc headings?
            #     left = width * row['Left']
            #     top = height * row['Top']
            #     #draw = ImageDraw.Draw(image)
            #     draw.rectangle([left, top, left + (width * row['Width']), top + (height * row['Height'])],
            #                    outline='pink')

            img_copy = image.copy()
            background = ImageDraw.Draw(img_copy, 'RGBA')
            background.rectangle([0, 0, image.size[0], image.size[1]], fill='green')
            image = Image.blend(img_copy, image, alpha=0.3)

        if report.fig_pages:
            if float(page[0]) in report.fig_pages['PageNum'].values: # change colour of fig pages
                img_copy = image.copy()
                background = ImageDraw.Draw(img_copy, 'RGBA')
                background.rectangle([0, 0, image.size[0], image.size[1]], fill='purple')
                image = Image.blend(image, img_copy, alpha=0.3)

        #else:
        # draw bb around section headers
        if int(page[0]) in report.section_ptrs['PageNum'].values:
            lnnums = report.section_ptrs.loc[report.section_ptrs['PageNum'] == int(page[0])]['LineNum']
            for line in lnnums.values:
                linenum = line - 1
                line = page[1][linenum]
                box = line['BoundingBox']
                left = width * box['Left']
                top = height * box['Top']
                draw.rectangle([left, top, left + (width * box['Width']), top + (height * box['Height'])], outline='blue')

        if int(page[0]) in report.subsection_ptrs['PageNum'].values:
            lnnums = report.subsection_ptrs.loc[report.subsection_ptrs['PageNum'] == int(page[0])]['LineNum']
            for line in lnnums.values:
                linenum = line - 1
                line = page[1][linenum]
                box = line['BoundingBox']
                left = width * box['Left']
                top = height * box['Top']
                draw.rectangle([left, top, left + (width * box['Width']), top + (height * box['Height'])],
                               outline='green')

        drawn_images.append(image)
    outfile = paths.get_report_name(report.docid, local_path=True, file_extension='_boxed.pdf', file_num=report.filenum)
    drawn_images[0].save(outfile, save_all=True, append_images=drawn_images[1:])

## Adds bookmarks to sections and sub-bookmarks to subsections of a document, and saves as a new file.
# If also running draw_report, run that first as bookmarks will not be saved in that file.
def bookmark_report(report, test=False):
    if len(report.docinfo.keys()) == 0:
        return
    if test:
        report_file = paths.get_report_name(report.docid, local_path=True, file_extension='_boxed.pdf', file_num=report.filenum)
    else:
        report_file = paths.get_report_name(report.docid, local_path=True, file_extension='.pdf', file_num=report.filenum)
    output = PdfFileWriter()
    input = PdfFileReader(open(report_file, 'rb')) #'../' +
    ptrs = report.headings_intext
    for page in input.pages:
        output.addPage(page)

    output.addBookmark('Title Page', 0, fit='/FitB')
    if report.toc_page:
        output.addBookmark('Table of Contents', report.toc_page-1, fit='/FitB')
    section = None
    for i, row in ptrs.iterrows():
        #page, line = row['PageNum'], row['LineNum']
        #lnbb = report.docinfo[page][line-1]['BoundingBox']
        if row['Heading'] == 1:
            section = output.addBookmark(row['Text'], row['PageNum']-1, fit='/FitB')
        elif row['Heading'] == 2:
            if section:
                output.addBookmark(row['Text'], row['PageNum']-1, parent=section, fit='/FitB')
            else:
                output.addBookmark(row['Text'], row['PageNum']-1, fit='/FitB')  # add as a heading if section doesn't exist

    refpg = output.getPage(0).mediaBox
    width, height = float(refpg[2]), float(refpg[3])

    # add links between toc lines and their intext section
    #self.headings_intext, self.subheadings, self.headings
    if report.toc_page:
        toc_headings = pd.concat([report.headings, report.subheadings])
        for i, row in report.headings_intext.iterrows():
            if row.MatchesHeading == 0:
                continue
            toc_h = toc_headings.loc[int(row.MatchesI)]
            toc_bb = report.line_dataset.loc[(report.line_dataset.PageNum == report.toc_page) &
                                             (report.line_dataset.LineNum == toc_h.LineNum)].iloc[0]
            left = width * toc_bb['Left']
            top = height * (1 - toc_bb['Top'])
            #rectangle = [left, top, left + (width * toc_bb['Width']), top + (height * toc_bb['Height'])]
            rectangle = [left, top, left + (width * toc_bb['Width']), top - (height * toc_bb['Height'])]
            output.addLink(report.toc_page-1, row.PageNum-1, rect=rectangle, fit='/FitB')  # creates link from toc heading to section page

    #outfile = settings.get_report_name(report.docid, local_path=True, file_extension='_bookmarked.pdf')
    outfile = paths.get_bookmarked_file(report.docid, test=test, filenum=report.filenum)
    print(outfile)
    rpath = outfile.rsplit('/', 1)[0]
    if not os.path.exists(rpath):
        os.mkdir(rpath)
    output.write(open(outfile, 'wb'))

## Saves report section text to Word document
def save_report_sections(report):
    if len(report.docinfo.keys()) == 0:
        return
    doc = docx.Document()
    for section in report.section_content:
        doc.add_heading(section['Heading'], 1)
        p = doc.add_paragraph()
        for line in section['Content']:
            p.add_run(line + '\n')
        doc.add_page_break()
    doc.save(paths.get_report_name(report.docid, local_path=True, file_extension='_sections.docx', file_num=report.filenum))

## Saves Report class information to a json file .
def report2json(report, test=False):
    if test:
        local = 'test'
    else:
        local = True
    with open(paths.get_report_name(report.docid, local_path=local, file_extension='.json', file_num=report.filenum), "w") as f:
        frozen = jsonpickle.encode(report)
        json.dump(frozen, f)

## Removes unwanted reports (qld mining journals, WELCOM) values from datasets
# QLD Mining Journals are unwanted in datasets as their pages contain two columns of text, so lines are reconstructed incorrectly
# Well Completion Reports are unwanted as most of the time they do not contain a table of contents
# These are two exclusions for report types the Bookmarker is expected to work for.
def sanitise_datasets():
    rtitle = 'QGMJ' 
    rtype = 'WELCOM'
    ref = pd.read_excel('C:/Users/andraszeka/Documents/gsq-boreholes/investigations/QDEX_metada_export.xlsx', dtype={'REPNO': int})
    bad = ref.loc[ref.RTITLE.str.contains(rtitle) | ref.RTYPE.str.contains(rtype)]
    bad_docids = bad.REPNO.values
    names = ['marginal_lines', 'toc', 'fig', 'heading_id_toc', 'heading_id_intext'] # page id and page extraction datasets don't have DocID attribute
    datasets = [paths.get_dataset_path(name) for name in names]
    for dataset in datasets:
        if os.path.exists(dataset):
            try:
                data = pd.read_csv(dataset, dtype={'DocID': int})
            except ValueError:
                data = pd.read_csv(dataset)
                data.dropna(subset=['DocID'], inplace=True)
                data.DocID = data.DocID.astype(int)
            prelen = data.shape[0]
            data = data.loc[~data.DocID.isin(bad_docids)]
            postlen = data.shape[0]
            data.to_csv(dataset, index=False)
            print('Removed ', str(prelen-postlen), ' bad values from ', dataset)



## Remove unwanted docid files from restructpageinfo
# See sanitise_datasets() for what is unwanted
def sanitise_files():
    rtitle = 'QGMJ'
    rtype = 'WELCOM'
    ref = pd.read_excel('C:/Users/andraszeka/OneDrive - ITP (Queensland Government)/gsq-boreholes/investigations/QDEX_metada_export.xlsx',
                        dtype={'REPNO': int})
    bad = ref.loc[ref.RTITLE.str.contains(rtitle) | ref.RTYPE.str.contains(rtype)]
    bad_docids = bad.REPNO.values
    removed = []
    ids = paths.get_files_from_path('restructpageinfo')
    lines_docs = paths.get_files_from_path('restructpageinfo', get_file_paths=True)
    for id, lines_doc in zip(ids, lines_docs):
        docid, filenum = id[0], id[1]
        if docid in bad_docids:
            if not os.path.exists('nottraining/restructpageinfo/'):
                os.makedirs('nottraining/restructpageinfo/')
            os.rename(lines_doc, paths.get_restructpageinfo_file(docid, local_path=True, training=False, file_num=filenum))
            removed.append([docid, filenum])
    print("Removed: ", len(removed), ", ", removed)

if __name__ == '__main__':
    #sanitise_datasets()
    sanitise_files()
    # transform document pages into dataset of pages for toc classification, classify pages, and isolate toc
    # from toc page, transform content into dataset of headings for heading identification, identify headings, and return headings and subheadings
    # test_reports = ['30320', '42688', '95183', '2984', '57418', '75738', '111200']
    # #reports = test_reports #['30320'] # '30320' #'24352', '24526', '26853', '28066', '28184','28882', '30281', '31681', '23508', ] #,'23732',

    # reports = [['92099', '1']]
    # test = False
    # for report in reports:
    #     docid = report[0]
    #     filenum=report[1]
    #     start = time.time()
    #     r = Report(docid, filenum)
    #     if test:
    #         draw_report(r)
    #     bookmark_report(r, test)
    #     save_report_sections(r)
    #     report2json(r, test=test)
    #     end = time.time()
    #     print('time:', end - start)

        #print('TOC Headings: \n')
        #for string in r.doclines[str(r.toc_page)]:
        #    print(string)
        #print("In Text Headings: \n", r.headings)
        #print("Subheadings: \n", r.subheadings)
        #print_sections(r)
